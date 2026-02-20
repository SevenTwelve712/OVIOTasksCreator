from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, Inches

from Configs import PathConfig
from src.utils.docx_documents_utils import do_rulang_xml

from docx.oxml import parse_xml

picture_xml = f"""<w:p {nsdecls('w', 'wp', 'r')}>
      <w:pPr>
        <w:pStyle w:val="Normal"/>
        <w:widowControl/>
        <w:bidi w:val="0"/>
        <w:spacing w:lineRule="auto" w:line="276" w:before="0" w:after="200"/>
        <w:jc w:val="left"/>
        <w:rPr/>
      </w:pPr>
      <w:r>
        <w:rPr/>
        <w:drawing>
          <wp:anchor behindDoc="0" distT="0" distB="0" distL="114300" distR="114300" simplePos="0" locked="0" layoutInCell="0" allowOverlap="1" relativeHeight="2">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
              <wp:posOffset>0</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
              <wp:posOffset>-56515</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="598170" cy="598170"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapTight wrapText="bothSides">
              <wp:wrapPolygon edited="0">
                <wp:start x="-1" y="0"/>
                <wp:lineTo x="-1" y="20642"/>
                <wp:lineTo x="20635" y="20642"/>
                <wp:lineTo x="20635" y="0"/>
                <wp:lineTo x="-1" y="0"/>
              </wp:wrapPolygon>
            </wp:wrapTight>
            <wp:docPr id="1" name="logo" descr="Olympiad_logo"/>
            <wp:cNvGraphicFramePr>
              <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
            </wp:cNvGraphicFramePr>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                  <pic:nvPicPr>
                    <pic:cNvPr id="0" name="logo.tif" descr="Olympiad_logo"/>
                    <pic:cNvPicPr>
                      <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
                    </pic:cNvPicPr>
                  </pic:nvPicPr>
                  <pic:blipFill>
                    <a:blip r:embed="rId9"/>
                    <a:stretch>
                      <a:fillRect/>
                    </a:stretch>
                  </pic:blipFill>
                  <pic:spPr bwMode="auto">
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="598170" cy="598170"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                    <a:noFill/>
                  </pic:spPr>
                </pic:pic>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </w:r>
    </w:p>"""


def do_template_docx():
    doc = Document()

    doc.add_picture(str(Path(PathConfig.RESOURCES_DIR, "logo.tif")), width=Inches(2))
    pic = doc._element.body.find(qn("w:p"))
    doc._element.body.remove(pic)

    picture = parse_xml(picture_xml)
    doc._element.body.insert(0, picture)

    section = doc.sections[0]

    doc.core_properties.language = "ru-RU"

    # Поля
    section.left_margin = Cm(1.27)
    section.top_margin = Cm(0.5)
    section.right_margin = Cm(1.27)
    section.bottom_margin = Cm(1.0)

    styles = []

    # Стили
    e_test_style = doc.styles.add_style('ETestTask', WD_STYLE_TYPE.PARAGRAPH)
    e_test_style.font.name = 'Times New Roman'
    e_test_style.font.size = Pt(10)
    styles.append(e_test_style)

    ovio_header_style = doc.styles.add_style('OVIOHeader', WD_STYLE_TYPE.PARAGRAPH)
    ovio_header_style.font.name = 'Times New Roman'
    ovio_header_style.font.size = Pt(12)
    styles.append(ovio_header_style)

    reading_style = doc.styles.add_style('ReadingTask', WD_STYLE_TYPE.PARAGRAPH)
    reading_style.base_style = doc.styles["No Spacing"]
    reading_style.font.name = 'Times New Roman'
    reading_style.font.size = Pt(10)
    styles.append(reading_style)



    # Заменяем язык на русский во всех стилях
    for style in styles:
        style._element.get_or_add_rPr().append(do_rulang_xml())

    doc.save(str(PathConfig.TEMPL_PATH))

def _get_basedoc_xml():
    print(Document(str(PathConfig.TEMPL_PATH))._element.xml)


if __name__ == "__main__":
    do_template_docx()
    _get_basedoc_xml()