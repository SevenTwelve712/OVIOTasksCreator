from docx.enum.section import WD_SECTION

from docx import Document
from pathlib import Path

from docx.oxml import OxmlElement, ns

from model.extended_docx_classes.ExtendedSection import ExtendedSection
from src.model.Task import Task
from src.model.TourTemplate import TourTemplate


class EruditionTest(Task):
    name = "Тест"

    def __init__(self, questions: dict[str: tuple[list[str], str]], template: TourTemplate):
        """
        Класс теста на эрудицию.
        :par am questions: Список вопросов в виде {"вопрос": (["4 варианта ответа"], 'буква правильного ответа')
        :param template: Шаблон заглавной плашки данного тура олимпиады
        """
        self.questions = questions
        self.template = template
        self.doc = None

    def make_xml(self):
        pass


    def make_docx(self, doc: Document):
        letters = ['А', 'Б', 'В', 'Г']
        self.template.make_docx(doc, self.name, "")
        doc = self.template.doc

        # Добавляем главную секцию с колонками
        main_sec = doc.add_section(WD_SECTION.CONTINUOUS)
        ExtendedSection(main_sec).set_cols(3)

        num = 0
        for question, answers_ in self.questions.items():
            num += 1
            par = doc.add_paragraph(style='ETestTask')
            answers, letter = answers_

            # Прогон вопроса
            question_run = par.add_run(f"{str(num)}. {question}")
            question_run.bold = True

            # Прогон вариантов ответов
            answer_run = par.add_run()

            for a_num, answer in enumerate(answers):
                answer_run.add_break()
                answer_run.add_text(f"{letters[a_num]}. {answer}")
        self.doc = doc