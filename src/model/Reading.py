from random import shuffle

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Emu

from .TourTemplate import TourTemplate
from .extended_docx_classes.ExtendedCell import ExtendedCell
from .extended_docx_classes.ExtendedParagraph import ExtendedParagraph, JcTypes
from .extended_docx_classes.ExtendedSection import ExtendedSection
from .extended_docx_classes.ExtendedTable import ExtendedTable
from .extended_docx_classes.TableHelpClasses import WidthTypes, TblBorder, LayoutTypes
from .extended_docx_classes.ExtendedRun import ExtendedRun
from .Task import Task


class ReadingTable1Task:
    def __init__(self, doc: Document, rows: int,  words: list[str], preferred_font_size: int, layout: LayoutTypes):
        """
        Реализует класс для удобной работы с объектом таблицы из первого задания для чтения. Она состоит из 1 строки и
        нескольких столбцов (по количеству слов), в каждой ячейке находится одно слово
        :param doc: Документ, в который должна быть добавлена таблица.
        :param words: Слова, которые должны находиться в таблице
        :param preferred_font_size: Предпочитаемый размер шрифта (может быть уменьшен)
        """
        if preferred_font_size not in [8, 10, 12]:
            raise ValueError("Размер шрифта должен быть 8, 10 или 12")

        self.ext_table = None
        self._table = None
        self.doc = doc
        self.words = words
        self.font_size = preferred_font_size
        self.default_table_width = self._calc_document_text_area_width()
        self.layout = layout
        self.rows = rows

    def create_table(self):
        """Добавляет саму таблицу в конец документа"""
        table = self.doc.add_table(self.rows, len(self.words))
        self.ext_table = ExtendedTable(table)
        self._table = table

    def _calc_document_text_area_width(self):
        """Считает ширину текстового поля документа (для того, чтобы заполнить таблицу на всю ширину), возвращает ответ в dxa"""
        emu_width = ExtendedSection(self.doc.sections[-1]).get_text_area_width()
        return Emu(emu_width).pt * 20

    def fill_words(self, bold: bool=False):
        """Заполняет таблицу заданными словами"""
        if self._table is None:
            raise RuntimeError("Таблица еще не создана")
        for i, cell in enumerate(self._table.rows[0].cells):
            run = cell.paragraphs[0].add_run()
            run.text = f"{i}. {self.words[i]}"
            run.bold = bold
            run.font.size = Pt(self.font_size)
            run.font.name = "Times New Roman"

    def set_preferred_grid_cols_widths(self):
        """Задает предпочитаемую ширину для колонок"""
        grid_widths = [Pt(self.calc_min_cell_width(word)).twips for word in self.words]
        self.ext_table.set_grids(grid_widths)

    def calc_min_cell_width(self, text: str) -> float:
        """Считает примерную максимальную достижимую ширину ячейки в dxa (twips), учитывая шрифт (Times New Roman) и его размер, цифра получается очень приблизительная,
        она получена эмпирическим опытом"""
        if self.font_size == 8:
            return (len(text) * 59 + 75) / 7 * 20
        elif self.font_size == 10:
            return (75 + 73 * len(text)) / 7 * 20
        else:
            return (75 + 87 * len(text)) / 7 * 20

    def normalize_widths(self):
        """Пытается нормализовать ширину таблицы: выставляет фиксированную ширину таблицы,
        задает предпочитаемые ширину колонок, устанавливает автоматическое распределение ширины"""
        if self.layout is LayoutTypes.AUTOFIT:
            self.ext_table.set_layout(LayoutTypes.AUTOFIT)
            self.ext_table.set_width(WidthTypes.DXA, width=self.default_table_width)
            self.set_preferred_grid_cols_widths()
        else:
            raise ValueError("Реализация для фиксированного layout еще не написана")


class Reading(Task):
    name = "Чтение"
    cond = "Познакомьтесь с текстом и выполните задания"

    def __init__(self, tour_templ: TourTemplate, text: str, matches: dict[str, str], questions: list[str], word: tuple[str, str], mistake_words: tuple[str, str]):
        """
        Класс задания чтение
        :param tour_templ: Шаблон заголовка задания.
        :param text: Текст, по которому выполняется задание.
        :param matches: Задание соответствий (задание 1), в формате {слово: ассоциация}.
        :param questions: Вопросы (задание 2).
        :param word: Слово по определению (задание 3), в формате (слово, его определение).
        :param mistake_words: Ошибочное и верные слова в формате (слово, слово)
        """

        self.text = text
        self.matches = matches
        self.questions = questions
        self.word = word
        self.mistake_words = mistake_words
        self.doc = None

        super().__init__(tour_templ)

    def make_xml(self):
        pass

    def make_docx(self, doc: Document):
        PG_MAR_HEADER = {"top": 250, "bottom": 250, "left": 720, "right": 720, "header": 708, "footer": 339, "gutter": 0}
        PG_MAR_TEXT = {"top": 250, "bottom": 250, "left": 400, "right": 400, "header": 708, "footer": 339, "gutter": 0}
        PG_MAR_TASKS = {"top": 250, "bottom": 250, "left": 720, "right": 720, "header": 708, "footer": 339, "gutter": 0}

        TEXT_INDENT = 100
        TASK_INDENT = 300

        SMALL_SPACING = Pt(3)

        PAR_STYLE = "ReadingTask"

        doc = super().make_docx(doc)

        ExtendedSection(doc.sections[0]).set_margins(*PG_MAR_HEADER.values())
        ExtendedParagraph(doc.paragraphs[0]).set_spacing(after=0)

        #========================================
        # Добавляем секцию текста
        text_sec = doc.add_section(WD_SECTION.CONTINUOUS)

        extended_first = ExtendedSection(text_sec)
        extended_first.set_cols(2, space=500)
        extended_first.set_size_a4()
        extended_first.set_margins(*PG_MAR_TEXT.values())

        text_par = doc.add_paragraph(self.text, style="ReadingTask")

        ext_text_par = ExtendedParagraph(text_par)
        ext_text_par.set_indent(left=TEXT_INDENT, first_line=107)
        ext_text_par.set_jc(JcTypes.BOTH)
        ext_text_par.rm_spacings()

        for run in text_par.runs:
            ExtendedRun(run).set_spacing(-2)

        # ========================================
        # добавляем секцию заданий
        tasks_sec = doc.add_section(WD_SECTION.CONTINUOUS)

        ExtendedSection(tasks_sec).set_cols(1)
        ExtendedSection(tasks_sec).set_size_a4()

        # ========================================
        # 1 задание
        f_task_par = doc.add_paragraph(style="ReadingTask")

        ext_f_task_par = ExtendedParagraph(f_task_par)
        ext_f_task_par.set_jc(JcTypes.BOTH)
        ext_f_task_par.set_indent(TASK_INDENT, TASK_INDENT, 0)
        ext_f_task_par.set_spacing(before=SMALL_SPACING, after=SMALL_SPACING)

        f_task_par.add_run("1. Заполните таблицу. Под каждым словом запишите НОМЕР соответствующего ему слова из списка (по 1 баллу за соответствие):").bold = True

        # Создаем таблицу со словами
        f_task_cond = doc.add_table(rows=1, cols=len(self.matches), style="Table Grid")
        for i, word, cell in zip(range(0, len(self.matches)), self.matches, f_task_cond.rows[0].cells):
            par = cell.paragraphs[0]

            ExtendedParagraph(par).set_spacing(after=SMALL_SPACING)
            ExtendedParagraph(par).set_jc(JcTypes.CENTER)
            par.style = PAR_STYLE

            par.add_run(f"{i + 1}. {word}")

        f_task_cond_ext = ExtendedTable(f_task_cond)
        f_task_cond_ext.set_borders(TblBorder())
        f_task_cond_ext.set_indent(TASK_INDENT)
        f_task_cond_ext.set_jc(JcTypes.CENTER)

        # Создаем таблицу с ассоциациями
        associations = list(self.matches.values())
        shuffle(associations)

        f_task_solution = doc.add_table(rows=2, cols=len(associations))
        for association, cell in zip(associations, f_task_solution.rows[0].cells):
            par = cell.paragraphs[0]

            ExtendedParagraph(par).set_jc(JcTypes.CENTER)
            par.style = PAR_STYLE

            cell.paragraphs[0].add_run(f"{association.upper()}")

        for cell in f_task_solution._cells:
            ExtendedParagraph(cell.paragraphs[0]).rm_spacings()

        tbl_borders = TblBorder(sz=4, val="single")
        f_task_solution_ext = ExtendedTable(f_task_solution)
        f_task_solution_ext.set_borders(tbl_borders)
        f_task_solution_ext.set_indent(TASK_INDENT)
        f_task_solution_ext.set_jc(JcTypes.CENTER)
        f_task_solution_ext.set_cell_spacing(15)
        f_task_solution_ext.set_all_cells_borders(tbl_borders)

        # ========================================
        # 2 задание
        S_TASK_ROWS = len(self.questions)
        S_TASK_COLS = 2

        s_task_par = doc.add_paragraph(style="ReadingTask")

        ext_s_task_par = ExtendedParagraph(s_task_par)
        ext_s_task_par.set_jc(JcTypes.BOTH)
        ext_s_task_par.set_indent(TASK_INDENT, TASK_INDENT, 0)
        ext_s_task_par.set_spacing(before=SMALL_SPACING, after=SMALL_SPACING)

        s_task_par.add_run("2. Заполните таблицу (по 2 балла за правильное заполнение. Слова должны быть написаны без ошибок):").bold = True

        s_task = doc.add_table(rows=S_TASK_ROWS, cols=S_TASK_COLS)
        dxa_width = Emu(ExtendedSection(doc.sections[-1]).get_text_area_width()).pt * 20
        print(dxa_width)
        grids = [int(dxa_width * 3 / 4), int(dxa_width / 4)]
        print(grids)

        s_task_ext = ExtendedTable(s_task)

        s_task_ext.set_layout(LayoutTypes.FIXED)
        s_task_ext.set_grids(grids)

        s_task_ext.set_borders(tbl_borders)
        s_task_ext.set_indent(TASK_INDENT)
        s_task_ext.set_jc(JcTypes.CENTER)
        s_task_ext.set_cell_spacing(15)

        for i, (question, cell) in enumerate(zip(self.questions, s_task.column_cells(0))):
            par = cell.paragraphs[0]
            par.add_run(f"2.{i}. {question}")
            par.style = PAR_STYLE

        for cell in s_task._cells:
            ExtendedCell(cell).set_width(WidthTypes.AUTO)

        s_task_ext.rm_spacings_in_cells()
        s_task_ext.set_all_cells_borders(tbl_borders)

        # ========================================
        # 3 задание
        t_task_par = doc.add_paragraph(style="ReadingTask")

        ext_t_task_par = ExtendedParagraph(t_task_par)
        ext_t_task_par.set_jc(JcTypes.BOTH)
        ext_t_task_par.set_indent(TASK_INDENT, TASK_INDENT, 0)
        ext_t_task_par.set_spacing(before=SMALL_SPACING, after=SMALL_SPACING)

        r = t_task_par.add_run("3. Определите слово по описанию (2 балла). Это слово обязательно должно быть в тексте.")
        r.bold = True
        r.add_break()
        t_task_par.add_run(f"{'_' * int(len(self.word[0]) / 0.7)} — {self.word[1]} ({len(self.word[0])} букв)")

        # ========================================
        # 4 задание
        fo_task_par = doc.add_paragraph(style="ReadingTask")

        ext_fo_task_par = ExtendedParagraph(fo_task_par)
        ext_fo_task_par.set_jc(JcTypes.BOTH)
        ext_fo_task_par.set_indent(TASK_INDENT, TASK_INDENT, 0)
        ext_fo_task_par.set_spacing(SMALL_SPACING, SMALL_SPACING)

        fo_task_par.add_run("4. Найдите в тексте ошибочное слово и замените его на верное (найденное – 1 балл, правильная замена – 1 балл):").bold = True

        fo_task = doc.add_table(rows=2, cols=2)

        ext_fo_task = ExtendedTable(fo_task)
        ext_fo_task.set_indent(TASK_INDENT)
        ext_fo_task.set_jc(JcTypes.CENTER)
        ext_fo_task.set_borders(tbl_borders)

        for word, cell in zip(("Ошибочное", f"Правильное ({len(self.mistake_words[1])} букв)"), fo_task.row_cells(0)):
            par = cell.paragraphs[0]
            par.style = PAR_STYLE
            par.add_run(word).bold = True
            ExtendedParagraph(par).set_jc(JcTypes.CENTER)

        ext_fo_task.rm_spacings_in_cells()

        self.doc = doc