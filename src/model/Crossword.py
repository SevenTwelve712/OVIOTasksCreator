from enum import EnumType
from xml.dom.minidom import Document

from docx.enum.section import WD_SECTION

from src.model.Task import Task
from src.model.TourTemplate import TourTemplate
from src.model.extended_docx_classes.ExtendedParagraph import ExtendedParagraph
from src.model.extended_docx_classes.ExtendedSection import ExtendedSection
from src.model.extended_docx_classes.data_and_enums import Direction, JcTypes
from src.model.vendor.complexstring import ComplexString
from src.model.vendor.genxword import Crossword

from io import BytesIO


class HeightTypes(EnumType):
    PX = "pixels"
    CELLS = "cells"


# TODO: test
class OVIOCrossword(Task):
    name = "Кроссворд"
    cond = ""

    CELL_SIZE = 30 # размер одной ячейки кроссворда в px для генерации jpg
    BORDER_SIZE = 5 # толщина границы вокруг каждой ячейки в px

    def __init__(self, words: list[tuple[str, str]], max_height: int, height_type: HeightTypes, tour_template: TourTemplate):
        """
        Класс задания кроссворд
        :param words: слова в формате списка из множеств (слово, описание)
        :param max_height: максимальная высота (в пикселях или в ячейках)
        :param height_type: тип максимальной высоты
        """
        super().__init__(tour_template)

        self.words = words

        if height_type is HeightTypes.PX:
            self.max_height = self._px_to_cells(max_height)
        else:
            self.max_height = max_height

        self._words_clues = [[ComplexString(word.upper()), clue] for word, clue in words]

    def _px_to_cells(self, px_size: int):
        return int(px_size / self.CELL_SIZE - self.BORDER_SIZE * 2)

    def make_docx(self, doc: Document):
        SECT_MAR = {"top": 250, "bottom": 250, "left": 720, "right": 720, "header": 708, "footer": 339,
                    "gutter": 0}
        CROSS_TIME_GENERATING = 0.02
        STYLE = "ReadingTask"
        doc = super().make_docx(doc)

        # =============================
        # Do img section
        img_sec = doc.add_section(WD_SECTION.CONTINUOUS)
        ExtendedSection(img_sec).set_size_a4()
        ExtendedSection(img_sec).set_margins(**SECT_MAR)

        # generating crossword
        COLS = self._px_to_cells(ExtendedSection(img_sec).get_text_area_width() / 9525)
        ROWS = self.max_height
        print(COLS, ROWS)
        cross = Crossword(ROWS, COLS, available_words=self._words_clues)
        cross.compute_crossword(CROSS_TIME_GENERATING)
        cross.remove_blank_lines()

        img = cross.gen_img(self.CELL_SIZE, self.BORDER_SIZE)
        img_bytes = BytesIO()
        img.save(img_bytes, format="JPeG")

        doc.add_picture(img_bytes)

        # =============================
        # Do clues section
        clues_sec = doc.add_section(WD_SECTION.CONTINUOUS)
        ExtendedSection(clues_sec).set_size_a4()
        ExtendedSection(clues_sec).set_margins(**SECT_MAR)
        ExtendedSection(clues_sec).set_cols(2)
        print(f"len(best_wordlist) is 16: {len(cross.best_wordlist) == 16}")

        for word in cross.best_wordlist:
            try:
                print(word[5], word[4])
            except IndexError:
                print(word)
                print("Error!!")
                raise IndexError
        words = sorted(cross.best_wordlist.copy(), key=lambda x: (x[5], x[4]))

        hor = doc.add_paragraph(style=STYLE)
        hor.add_run("По горизонтали:").bold = True
        ExtendedParagraph(hor).set_borders(24, 1, "#000000", Direction.BOTTOM)

        hor_clues = doc.add_paragraph(style=STYLE)
        ExtendedParagraph(hor_clues).set_jc(JcTypes.BOTH)
        for word in words:
            word, clue, x, y, align, num = tuple(word)
            if align == 1: # vertical words
                continue
            r = hor_clues.add_run(f"{num}. {clue}")
            r.add_break()

        ver = doc.add_paragraph(style=STYLE)
        ver.add_run("По вертикали:").bold = True
        ExtendedParagraph(hor).set_borders(24, 1, "#000000", Direction.BOTTOM)

        ver_clues = doc.add_paragraph(style=STYLE)
        ExtendedParagraph(hor_clues).set_jc(JcTypes.BOTH)
        for word in words:
            word, clue, x, y, align, num = tuple(word)
            if align == 0:  # horizontal words
                continue
            r = ver_clues.add_run(f"{num}. {clue}")
            r.add_break()

        self.doc = doc
