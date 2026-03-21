from docx.document import Document
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .extended_docx_classes.ExtendedSection import ExtendedSection
from .TourData import Forms, Tours


class TourTemplate:
    months = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря"
    }

    def __init__(self, form: Forms, tour: Tours, tour_date: date, place: str):
        self.tour = tour
        self.form = form
        self.tour_date = tour_date
        self.place = place
        self.doc = None

    def make_xml(self):
        pass

    def make_docx(self, doc: Document, task_name: str, task_cond: str):
        # TODO: разобраться, тут ли надо проставлять размер секции
        ExtendedSection(doc.sections[0]).set_size_a4()
        if len(doc.paragraphs) == 0:
            doc.add_paragraph()

        par = doc.paragraphs[0]
        par.style = "OVIOHeader"
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header = par.add_run("Открытая всероссийская интеллектуальная олимпиада «Наше наследие»")
        header.bold = True
        header.add_break()
        par.add_run(f"{self.tour.value} {self.tour_date.year}/{self.tour_date.year + 1} ({self.form.value} классы)").add_break()


        # datetime_data = par.add_run(f"{self.tour_date.day} {self.months[self.tour_date.month]} {self.tour_date.year}, {self.place}")
        # datetime_data.italic = True
        # datetime_data.add_break()

        par.add_run(f'ЗАДАНИЕ "{task_name}".').font.size = Pt(10)
        cond = par.add_run(task_cond)
        cond.bold = True
        cond.font.size = Pt(10)
        # cond.add_break()
        self.doc = doc
