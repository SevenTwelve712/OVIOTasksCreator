from pathlib import Path

from docx import Document
from docx.shared import Emu

from Configs import PathConfig
from src.model.Reading import ReadingTable1Task

from random import randint

RU_WORDS = [
    "дом", "лес", "кот", "сад", "мир", "стол", "шар", "нос", "сын", "день",
    "ночь", "утро", "окно", "поле", "река", "гора", "море", "ветер", "дождь", "снег",
    "огонь", "камень", "город", "улица", "школа", "книга", "ручка", "тетрадь", "учитель", "ученик",
    "машина", "поезд", "самолет", "корабль", "дорога", "мост", "здание", "комната", "квартира", "кухня",
    "работа", "зарплата", "профессия", "компания", "директор", "менеджер", "клиент", "проект", "задача", "решение",
    "компьютер", "ноутбук", "телефон", "экран", "клавиатура", "программа", "система", "файл", "документ", "таблица",
    "человек", "женщина", "мужчина", "ребенок", "девочка", "мальчик", "семья", "родитель", "друг", "подруга",
    "еда", "вода", "чай", "кофе", "суп", "хлеб", "сыр", "яблоко", "банан", "овощ",
    "музыка", "песня", "фильм", "театр", "актер", "сцена", "картина", "выставка", "музей", "искусство",
    "спорт", "игра", "команда", "матч", "тренер", "стадион"
]

def choose_random_word():
    word = RU_WORDS[randint(0, len(RU_WORDS) - 1)]
    if len(word) < 0:
        word = choose_random_word()
    return word

def test_ReadingTable1Task():
    TABLES_AMOUNT = 10
    COLS_AMOUNT = 6
    FONT_SIZE = 12
    SAVE_PATH = str(Path(PathConfig.SAVE_DIR, "ReadingTable1TaskTest.docx"))
    doc = Document()

    for _ in range(TABLES_AMOUNT):
        words = [choose_random_word() for __ in range(COLS_AMOUNT)]
        read_table = ReadingTable1Task(doc, words, FONT_SIZE)
        read_table.create_table()
        read_table.fill_words(bold=True)
        read_table.normalize_widths()
        doc.add_paragraph().add_run(" ")

    doc.save(SAVE_PATH)
    # doc = Document(SAVE_PATH)
    # table_xml = doc.tables[0]._element.xml
    # print(table_xml)

def get_test_ReadingTable1Task_xml():
    SAVE_PATH = str(Path(PathConfig.SAVE_DIR, "ReadingTable1TaskTest.docx"))
    doc = Document(SAVE_PATH)
    table_xml = doc.tables[0]._element.xml
    print(table_xml)

if __name__ == '__main__':
    test_ReadingTable1Task()
    # get_test_ReadingTable1Task_xml()