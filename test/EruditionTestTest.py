from datetime import date
from Configs import PathConfig
from pathlib import Path

from src.model.TourData import *

from docx import Document

from src.model.EruditionTest import EruditionTest
from src.model.TourTemplate import TourTemplate

def test_erudition_test():

    doc = Document(PathConfig.TEMPL_PATH)
    test_task = EruditionTest(
        {
            "В каком из перечисленных озер водится омуль?": (["В озере Байкал", "В Ладожском озере", "В озере Селигер", "В Телецком озере"], 'A'),
            "Самая крупная птица, обитающая на территории России:": (["Кайра", "Кудрявый пеликан", "Орлан-белохвост", "Черный гриф"], 'A'),
            "В каком из перечисленных храмов сохранились фрески иконописца Дионисия?": (["Воскресенский собор в Череповце", "Собор Рождества Богородицы в Ферапонтово", "Храм во имя святителя Николая Чудотворца в Вологде", "Храм Казанского образа Божией Матери в Устюжне"], 'A'),
            "Слово в списке, которое помимо своего основного значения является названием российской реки:": (["Дядя", "Мама", "Папа", "Тётя"], 'A'),
            "Какой из перечисленных крупных городов является ближайшим к указателю географического центра России, установленному на базальтовом плато Путорана?": (["Анадырь", "Воркута", "Магадан", "Норильск"], 'A'),
            "Самая крупная по площади в России заболоченная местность, по площади чуть меньшая Хорватии, расположенная в Томской, Омской и Новосибирской областях между реками Обь и Иртыш:": (["Васюганские болота", "Мшинское болото", "Сестрорецкое болото", "Старосельский мох"], 'A'),
            "Денежной единицей Древней Руси была гривна. По-древнерусски, считая деньги, можно было сказать, например, так: пол пяты гривны. Какое количество денег обозначалось таким способом?": (["Две с половиной гривны", "Одна десятая гривны", "Пять с половиной гривен", "Четыре с половиной гривны"], 'A')
        },
        TourTemplate(Forms.OLD, Tours.FINAL, date(2023, 11, 4), "Москва")
    )
    doc = test_task.make_docx(doc)
    last_p = len(doc.paragraphs) - 1

    print(f"Количество секций: {len(doc.sections)}")
    print(f"Количество параграфов: {len(doc.paragraphs)}")
    print(f"Текст последнего параграфа: {doc.paragraphs[last_p].text}")
    print(f"Стиль последнего параграфа: {doc.paragraphs[last_p].style}")

    xml_path = Path(PathConfig.BASE_DIR, "test", "result_files", "ETestTest.xml")
    docx_path = Path(PathConfig.BASE_DIR, "test", "result_files", "ETestTest.docx")
    with open(xml_path, 'w', encoding="utf-8") as f:
        print(doc._element.xml, file=f)
    doc.save(docx_path)

if __name__ == "__main__":
    test_erudition_test()
